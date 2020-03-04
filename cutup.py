import docx
import random
import sys
import getopt

#return 1: option value pair for deciding whether to use cutup or foldin, and 2: a list of files to mix, with the last one being output.
def getargs(arglist):
    args = getopt.getopt(arglist, "c:f:h")
    options = args[0]
    if len(options) != 1:
        print("Can only specify -c OR -f OR -h, not any or none", file=sys.stderr)
        return 0
    options = options[0]
    if options[0] == 'h':
        print(
        """Usage: 
        -c num implements the 'cutup' method, whereby Burroughs would randomly cut pages and shuffle them about. num here signifies the max length possible for a cut. 
        -f num implements the 'foldin' method, whereby Burroughs would fold two pages in half and layer them down the middle. num here signifies how many words to read from one text before 'folding in' the next.
        In both cases, specify at least 3 .docx files: 2 to mix together, and the 3rd as output. The output file need not already exist.
        You can mix as many files as you like, however the output will be as long as the shortest"""
        )
        return 0
    try:
        int(options[1])
    except:
        print("Length option must be integer", file=sys.stderr)
        return 0
    files = args[1]
    if len(files) < 3:
        print("Requires at least 3 non-option arguments: 2 files to mix together, a 3rd as output.", file=sys.stderr)          
        return 0
    for i in files:
        print(i)
    for i in files:
        if not i.endswith(".docx", -5):
            print("All non-option arguments must be .docx files", sys.stderr)
            return 0
    return (options, files)

def scanfile(path):     #converts a path ending in a .docx into a string of its content
    file = docx.Document(path)
    content = []
    for paragraph in file.paragraphs:
        content.append(paragraph.text)
    return(''.join(content))

def stringsplit(stringlist):    #convert a list of strings to a list of lists of words
    splitstrings = []
    for i in stringlist:
        splitstrings.append(i.split())
    for i in splitstrings:
        instindx = range(1, len(i), 200)
        for j in instindx:
            i.insert(j, '\n')
    return splitstrings
 
#produces a string of randomly sized chunks (up to size length) alternating between the two texts
def cutup(filelist, length):
    #necessary, because chunks will be a list of sentences: need a separate list to keep track of word count over sentence count
    #this is itself necessary, because when we call shuffle on chunks we want to shuffle sentences not words
    chunks = []
    wcount = []
    size = len(min(filelist, key=len))
    a = 1; b = 1
    #stop once size of shortest document reached
    while len(wcount) < size:
        c = random.randrange(1, length)
        b += c
        for i in filelist:
            chunk = i[a:b]
            wcount.extend(chunk)
            chunks.append(' '.join(chunk))
        a = b
    random.shuffle(chunks)
    return ' '.join(chunks)

def foldin(filelist, length): #alternates chunk by chunk through both, maintaining the order. Will produce the same result every time given the same parameters
    chunks = []
    size = len(min(filelist, key=len))
    a = 1; b = a + length
    while len(chunks) < size:
        for i in filelist:
            chunk = i[a:b]
            chunks.extend(chunk)
            a = b
            b += length
    return ' '.join(chunks)

def choose(filelist, savedir, length=None, cut=False):
    if cut:
        content = cutup(filelist, length)
        final = docx.Document()
        final.add_paragraph(content)
        final.save(savedir)
    else:
        content = foldin(filelist, length)
        final = docx.Document()
        final.add_paragraph(content)
        final.save(savedir)
    
def main():
    paths = getargs(sys.argv[1:])
    if not paths:
        return
    elif paths[0][0] == '-c':
        x = True
    else:
        x = False
    documents = []
    for i in paths[1][:-1]:
        documents.append(scanfile(i))
    documents = stringsplit(documents)
    choose(documents, paths[1][-1], length=int(paths[0][1]), cut=x)
    print("Done, output saved at %s" % paths[1][-1])

    


if __name__ == '__main__':
    main()


#needs: dialog to choose save location and title
#choosing the start point of fold in method i.e. for a longer second text it can be started in the middle
#specify the length of the final product
#option to shuffle or not

