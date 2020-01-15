import docx
from pathlib import Path
import random
    
def doc_select():       #navigates to and returns a file path
    path = Path().home()
    print('choose a .docx file. simply press enter to go up a level')
    input()
    while True:
        ls = []
        print('****', path.name, '****')
        for i in path.iterdir():
            print(i.name)
            ls.append(i.name.lower())   #prints all objects in directory
        while True: #checks if your input is in the directory or if you wish to go up a level
            x = input().lower()
            if x in ls:
                path = str(path) + '\\' + x
                path = Path(path)
                break
            elif x == '':
                path = str(path)
                path = path[:path.rfind('\\')]
                path = Path(path)
                break
            else: print('object not found')
        y = path.is_dir()
        if y == False:
            if str(path).endswith('.docx'): break
            else:
                print('needs to be .docx file')
                path = str(path)
                path = path[:path.rfind('\\')]
                path = Path(path)
    return path

def fileassignment():       #creates a list of paths of each file you want to mash up
    fpaths = []
    while True:
        fpaths.append(doc_select())
        if len(fpaths) > 1:
            x = input('{} files selected. press enter to add another, or any other input to continue'.format(len(fpaths)))
            if x == '': continue
            else: return fpaths                    


def scanfile(path):     #converts a path ending in a .docx into a string of its content
    file = docx.Document(path)
    content = []
    for paragraph in file.paragraphs:
        content.append(paragraph.text)
    return(''.join(content))

def pathtotext(fpaths):     #converts a list of paths to a list of strings
    text = []
    for i in fpaths:
        text.append(scanfile(i))
    return text

def stringsplit(stringlist):    #convert a list of strings to a list of lists of words
    splitstrings = []
    for i in stringlist:
        splitstrings.append(i.split())
    for i in splitstrings:
        instindx = range(1, len(i), 200)
        for j in instindx:
            i.insert(j, '\n')
    return splitstrings
        


def cutup(filelist, length=None, rdm=False, lim=False): #takes a list of strings and mashes up into one

    def randomise(filelist, length, lim):
        if length == None:  # this tends to basically split the two in half.. doesnt work too well!
            chunks = []
            wcount = []
            size = len(min(filelist, key=len))
            while len(wcount) < size:
                for i in filelist:
                    a = 2; b = 1
                    while a > b:
                        a = random.randrange(1, len(i)); b = random.randrange(1, len(i)) 
                    chunk = i[a:b]
                    del i[a:b]
                    wcount.extend(chunk)
                    chunks.append(' '.join(chunk))           
            return ' '.join(chunks)      
        elif length:
            if lim:     #produces a string of randomly sized chunks (up to length) alternating between the two texts
                chunks = []
                wcount = []
                size = len(min(filelist, key=len))
                a = 1; b = 1
                while len(wcount) < size:
                    c = random.randrange(1, length)
                    b += c
                    for i in filelist:
                        chunk = i[a:b]
                        wcount.extend(chunk)
                        chunks.append(' '.join(chunk))
                    a = b
                random.shuffle(chunks)
                return ' '.join(chunks)
            else:       #produces a string of chunks (of random but same size, determined by length param) alternating between the two texts
                chunks = []
                wcount = []
                size = len(min(filelist, key=len))
                a = 1; c = random.randrange(1, length); b = a + c
                while len(wcount) < size:
                    for i in filelist:
                        chunk = i[a:b]
                        wcount.extend(chunk)
                        chunks.append(' '.join(chunk))
                    a = b
                    b += c
                random.shuffle(chunks)
                return ' '.join(chunks)

    def foldin(filelist, length): #alternates chunk by chunk through both, maintaining the order
        chunks = []
        size = len(min(filelist, key=len))
        a = 1; b = a + length
        while len(chunks) < size:
            for i in filelist:
                chunk = i[a:b]
                chunks.extend(chunk)
                a = b
                b += length
        return ' '.join(chunks)


        
    if rdm:
        content = randomise(filelist, length, lim)
        final = docx.Document()
        final.add_paragraph(content)
        final.save(Path(str(Path().home()) + '\\output.docx'))  #need to add code for choosing save dir

    elif not rdm:
        content = foldin(filelist, length)
        final = docx.Document()
        final.add_paragraph(content)
        final.save(Path(str(Path().home()) + '\\output.docx'))

    #cutup needs to return something, string of file path?
    
def main():
    paths = fileassignment()
    content = pathtotext(paths)
    content = stringsplit(content)
    while True:
        mode = input('''at what point do you want the lines to be cut?
for example, there are normally 20 words per line on a standard a4 page,
so enter "10" if you want the lines cut in half,
or "20" if you want to roughly alternate line by line.
otherwise, type "random", for a "cut-up" as opposed to "fold-in" method.
if using "random" mode, optional arguments are -lim and an integer.
-lim means the following integer is the limit for random selection of word length
other wise each chunk will be the length of said integer\n''')
        if mode[:6] == 'random':
            l = False
            r = True
            length = None
            if mode[7:11] == '-lim':
                l = True
                try:
                    length = int(mode[12:])
                    break
                except:
                    print('must be in format random -lim integer. -lim is optional')
            elif len(mode) > 6:
                 try:
                     length = int(mode[7:])
                     break
                 except:
                    print('must be in format random -lim integer. -lim is optional')
            else: break
        else:
            l = False
            r = False
            try:
                length = int(mode)
                break
            except:
                print('for fold-in mode argument must be integer')
    cutup(content, length, rdm=r, lim=l)

if __name__ == '__main__':
    main()


#needs: dialog to choose save location and title
#choosing the start point of fold in method i.e. for a longer second text it can be started in the middle
#specify the length of the final product
#option to shuffle or not

