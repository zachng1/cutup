import docx
import random
from pydub import AudioSegment

def scanfile(path):
    """
    Return a docx file's contents as a string.
    """
    file = docx.Document(path)
    content = []
    for paragraph in file.paragraphs:
        content.append(paragraph.text)
        content.append('\n')
    return(''.join(content))

def stringsplit(stringlist):
    """
    Takes a list of space separated strings i.e. ["word1", "word2 word3 word4", "word5 word6"] 
    and converts into a list of strings i.e. ["word1", "word2", "word3"..."wordn"]
    """
    splitstrings = []
    for i in stringlist:
        splitstrings.append(i.split())
    # for i in splitstrings:
    #     instindx = range(1, len(i), 200)
    #     for j in instindx:
    #         i.insert(j, '\n')
    return splitstrings
 
#produces a string of randomly sized chunks (up to size length) alternating between the two texts
def cutup(wordlist, length):
    chunks = []
    sizemax = len(max(wordlist, key=len))
    sizemin = len(min(wordlist, key=len))
    a = 0; b = 0

    #use size max and sizemin in this loop to ensure an even distribution of segments
    #also used so final output is only as long as the shortest
    while b < sizemax:
        c = random.randrange(1, length)
        b += c
        for i in wordlist:
            chunk = i[a:b]
            chunks.append(chunk)
        a = b
    random.shuffle(chunks)
    chunks = chunks[:sizemin]
    returnString = ""
    for i in chunks:
        if not i:
            continue
        returnString = returnString + ' '.join(i) + ' '
    return returnString

def foldin(wordlist, length): #alternates chunk by chunk through both, maintaining the order. Will produce the same result every time given the same parameters
    chunks = []
    size = len(min(wordlist, key=len))
    a = 0; b = a + length
    while len(chunks) < size:
        for i in wordlist:
            chunk = i[a:b]
            chunks.extend(chunk)
            a = b
            b += length
    return ' '.join(chunks)

def selectmode(wordlist, savedir, length, cut):
    if cut:
        print("Cutting")
        content = cutup(wordlist, length)
        final = docx.Document()
        final.add_paragraph(content)
        final.save(savedir)
    else:
        print("Folding")
        content = foldin(wordlist, length)
        final = docx.Document()
        final.add_paragraph(content)
        final.save(savedir)

